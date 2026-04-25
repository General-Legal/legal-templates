"""
Convert .docx legal templates to clean, LLM-optimized markdown.

Reads .docx files from docx-originals/ and writes markdown to
templates/<slug>/template.md.

Usage:
    python3 scripts/convert_docx.py

Requires:
    pip install python-docx

Design principles:
- Highlighted text becomes <mark>text</mark> (template placeholders).
- No programmatic templating syntax. The output is a readable legal
  document, not code.
- Preserve document structure: headings, numbered lists, tables,
  and paragraphs.
- Normalize non-ASCII whitespace and smart punctuation to ASCII
  equivalents for maximum compatibility.
- Signature-block tables are simplified into plain text.
"""

import re
from pathlib import Path
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

# Word XML namespace prefix
WML = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# Maps each .docx stem to the directory slug under templates/
TEMPLATE_MAP = {
    "Business Associate Agreement (BAA)": "business-associate-agreement",
    "Cookie Notice": "cookie-notice",
    "Mutual Non-Disclosure Agreement (NDA)": "mutual-nda",
    "One-Way Non-Disclosure Agreement (NDA)": "one-way-nda",
    "Privacy Policy (GDPR)": "privacy-policy-gdpr",
    "Privacy Policy (US)": "privacy-policy-us",
    "Terms of Use": "terms-of-use",
}


def is_highlighted(run):
    """Return True if the run has any highlight color applied."""
    hl = run.font.highlight_color
    return hl is not None and hl != WD_COLOR_INDEX.AUTO


def run_to_text(run):
    """Convert a single run to markdown text.

    Highlighted runs become <mark>...</mark>. Bold and italic are
    converted to standard markdown emphasis markers.
    """
    text = run.text
    if not text:
        return ""
    if is_highlighted(run):
        return f"<mark>{text}</mark>"
    if run.bold and run.italic:
        return f"***{text}***"
    if run.bold:
        return f"**{text}**"
    if run.italic:
        return f"*{text}*"
    return text


def para_to_text(para):
    """Assemble all runs in a paragraph into a single markdown string.

    Adjacent <mark> tags and bold markers are merged so the output
    doesn't contain empty or redundant formatting spans.
    """
    parts = [run_to_text(r) for r in para.runs]
    text = "".join(parts)
    # Merge adjacent marks: </mark><mark> -> empty
    text = re.sub(r"</mark>\s*<mark>", "", text)
    # Merge adjacent bold: **** -> empty
    text = re.sub(r"\*\*\*\*", "", text)
    return text.strip()


def get_list_level(para):
    """Determine the list nesting level of a paragraph.

    Returns (level, True) for list items, or None for non-list
    paragraphs. The level is 0-based (0 = top level).
    """
    style = para.style.name
    if style == "Tabbed_L1":
        return 0, True
    if style == "Tabbed_L2":
        return 1, True

    if style != "List Paragraph":
        return None

    numPr = para._element.find(f".//{WML}numPr")
    if numPr is None:
        return None

    level = 0
    ilvl_elem = numPr.find(f"{WML}ilvl")
    if ilvl_elem is not None:
        val = ilvl_elem.get(f"{WML}val")
        if val is not None:
            level = int(val)

    return level, True


def table_to_markdown(table):
    """Convert a docx table to markdown.

    Regular tables become pipe-delimited markdown tables. Signature
    block tables (detected by the presence of fields like "By:" and
    party labels like "Company") are converted to plain-text blocks
    instead, since their grid layout is not meaningful.
    """
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_text = cell.text.strip().replace("\n", " ")
            cells.append(cell_text)
        rows.append(cells)

    if not rows:
        return ""

    # Detect signature blocks
    all_text = " ".join(" ".join(r) for r in rows).lower()
    sig_labels = ["by:", "name:", "printed name:", "title:", "signature"]
    party_labels = ["company", "recipient", "signatory", "customer"]
    if any(s in all_text for s in sig_labels) and any(p in all_text for p in party_labels):
        return signature_block(rows)

    # Regular markdown table
    lines = []
    lines.append("| " + " | ".join(rows[0]) + " |")
    lines.append("| " + " | ".join("---" for _ in rows[0]) + " |")
    for row in rows[1:]:
        while len(row) < len(rows[0]):
            row.append("")
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def signature_block(rows):
    """Convert a signature table into a readable plain-text block.

    Merged cells in docx tables cause python-docx to repeat cell
    values, so we deduplicate before rendering. Party headers are
    bolded, and field labels are rendered as a simple list.
    """
    seen = []
    for row in rows:
        for cell in row:
            text = cell.strip()
            if text and text not in seen:
                seen.append(text)

    lines = ["", "---", "**Signature Block**", ""]
    for cell in seen:
        upper = cell.upper().replace(":", "").strip()
        if upper in ("COMPANY", "CUSTOMER", "RECIPIENT", "OTHER SIGNATORY"):
            lines.append(f"**{cell.strip().rstrip(':')}**")
            continue
        if ":" in cell and len(cell) < 100:
            lines.append(f"- {cell.strip()}")
            continue
        clean = cell.replace("_", "").strip()
        if clean:
            lines.append(f"- {clean}")
    lines.extend(["", "---", ""])
    return "\n".join(lines)


def normalize_text(text):
    """Replace non-ASCII whitespace and smart punctuation with ASCII.

    Preserves meaningful symbols like section (S) and copyright (c).
    """
    text = text.replace("\u00a0", " ")   # non-breaking space
    text = text.replace("\u2018", "'")    # left single quote
    text = text.replace("\u2019", "'")    # right single quote
    text = text.replace("\u201c", '"')    # left double quote
    text = text.replace("\u201d", '"')    # right double quote
    text = text.replace("\u2013", "-")    # en dash
    text = text.replace("\u2014", "--")   # em dash
    return text


def convert(docx_path):
    """Convert a .docx file to LLM-optimized markdown.

    Walks the document body in order, handling paragraphs, headings,
    numbered lists, and tables. Returns the full markdown string.
    """
    doc = Document(docx_path)
    lines = []
    prev_blank = False
    counters = {}

    for block in doc.element.body:
        tag = block.tag.split("}")[-1]

        if tag == "tbl":
            for t in doc.tables:
                if t._element is block:
                    md = table_to_markdown(t)
                    if md:
                        lines.append("")
                        lines.append(md)
                        lines.append("")
                        prev_blank = True
                    break
            continue

        if tag != "p":
            continue

        para = None
        for p in doc.paragraphs:
            if p._element is block:
                para = p
                break
        if para is None:
            continue

        text = para_to_text(para)
        if not text:
            if not prev_blank:
                prev_blank = True
            continue

        style = para.style.name

        # Headings
        if style.startswith("Heading"):
            try:
                level = int(style.split()[-1])
            except ValueError:
                level = 1
            lines.append("")
            lines.append(f"{'#' * level} {text}")
            lines.append("")
            prev_blank = True
            counters = {}
            continue

        # List items
        list_info = get_list_level(para)
        if list_info is not None:
            level, _ = list_info
            indent = "    " * level
            key = level
            counters[key] = counters.get(key, 0) + 1
            for k in list(counters):
                if k > level:
                    del counters[k]
            lines.append(f"{indent}{counters[key]}. {text}")
            prev_blank = False
            continue

        # Normal paragraph
        if not prev_blank and lines:
            lines.append("")
        lines.append(text)
        prev_blank = False
        counters = {}

    result = "\n".join(lines).strip()
    result = re.sub(r"\n{3,}", "\n\n", result)
    result = normalize_text(result)
    return result


def main():
    root = Path(__file__).resolve().parent.parent
    originals = root / "docx-originals"

    if not originals.exists():
        print(f"Error: {originals} not found.")
        return

    converted = 0
    for f in sorted(originals.glob("*.docx")):
        if f.name.startswith("~$"):
            continue
        slug = TEMPLATE_MAP.get(f.stem)
        if not slug:
            print(f"Skipping (no mapping): {f.name}")
            continue
        print(f"Converting: {f.name}")
        md = convert(f)
        dest = root / "templates" / slug / "template.md"
        dest.parent.mkdir(parents=True, exist_ok=True)
        dest.write_text(md + "\n", encoding="utf-8")
        print(f"  -> {dest.relative_to(root)} ({len(md):,} chars)")
        converted += 1

    print(f"\nConverted {converted} templates.")


if __name__ == "__main__":
    main()
