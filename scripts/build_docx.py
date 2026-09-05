"""
Compile the markdown templates back into .docx files.

Reads templates/<slug>/template.md and writes build/docx/<slug>.docx.
This is the inverse of convert_docx.py: markdown is the source of truth,
and the .docx files it produces are the distributable artifacts attached
to each release.

Usage:
    python3 scripts/build_docx.py                    # build every template
    python3 scripts/build_docx.py mutual-nda         # build one or more
    python3 scripts/build_docx.py --out dist --zip   # custom dir, plus a zip

Requires:
    pip install -r scripts/requirements.txt

Design principles:
- <mark>text</mark> becomes highlighted text, so a generated .docx
  round-trips through convert_docx.py with its placeholders intact.
  Every build re-opens its own output and compares the highlighted text
  against the source, and fails if they disagree. --no-check skips it.
- The output mirrors the markdown. Numbers on list items are rendered
  literally rather than handed to Word's list engine, so a .docx never
  silently disagrees with the template.md it came from. Pass --renumber
  if you want them resequenced instead.
- Formatting is deliberately plain: letter paper, one-inch margins,
  Times New Roman. These are documents people redline, not brochures.
- A few templates carry unbalanced emphasis markers from the original
  .docx conversion. Those are rendered literally and reported at the end
  of a run; they are defects in template.md, not in this script.
"""

import argparse
import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(11)
HIGHLIGHT = WD_COLOR_INDEX.YELLOW

# Matches a run of inline markup. Order matters: <mark> first so its
# contents are handled as a unit, then *** before ** before *.
INLINE_RE = re.compile(
    r"(<mark>.*?</mark>|\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*)",
    re.DOTALL,
)
HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
ORDERED_RE = re.compile(r"^(\s*)(\d+)\.\s+(.*)$")
BULLET_RE = re.compile(r"^(\s*)-\s+(.*)$")
TABLE_SEP_RE = re.compile(r"^\|[\s:|-]+\|$")
# A thematic break: three or more -, * or _ and nothing else. The offer
# letter uses a tab-separated "*\t*\t*" dinkus as a section divider.
THEMATIC_RE = re.compile(r"^\s*(?:(?:-\s*){3,}|(?:\*\s*){3,}|(?:_\s*){3,})$")


# --------------------------------------------------------------------------
# Markdown -> block list
# --------------------------------------------------------------------------

def split_table_row(line):
    """Split a pipe-delimited row into cells, dropping the outer pipes."""
    return [c.strip() for c in line.strip().strip("|").split("|")]


def parse_blocks(md):
    """Turn markdown into a flat list of (kind, payload) blocks.

    Kinds: heading, rule, table, ordered, bullet, paragraph. Everything
    the converter emits is covered; anything unrecognized falls through
    to a paragraph so no text is ever dropped.
    """
    lines = md.replace("\r\n", "\n").split("\n")
    blocks = []
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if THEMATIC_RE.match(line):
            blocks.append(("rule", None))
            i += 1
            continue

        heading = HEADING_RE.match(stripped)
        if heading:
            blocks.append(("heading", (len(heading.group(1)), heading.group(2).strip())))
            i += 1
            continue

        # A table is a run of consecutive pipe rows. The |---|---| row is
        # a separator, not data.
        if stripped.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = lines[i].strip()
                if not TABLE_SEP_RE.match(row):
                    rows.append(split_table_row(row))
                i += 1
            if rows:
                blocks.append(("table", rows))
            continue

        ordered = ORDERED_RE.match(line)
        if ordered:
            indent, number, text = ordered.groups()
            blocks.append(("ordered", (len(indent) // 4, number, text.strip())))
            i += 1
            continue

        bullet = BULLET_RE.match(line)
        if bullet:
            indent, text = bullet.groups()
            blocks.append(("bullet", (len(indent) // 2, text.strip())))
            i += 1
            continue

        # A bold or italic span occasionally straddles a line break --
        # "**Annex 1" on one line and "**" on the next. Markdown would
        # join them; this parser is line-oriented because the templates
        # otherwise use one line per paragraph (address blocks, index
        # entries) and joining everything would run them together. So
        # join only when doing so resolves a dangling marker.
        if has_stray_marker(stripped):
            merged, consumed = stripped, 0
            for lookahead in lines[i + 1:i + 3]:
                if not lookahead.strip():
                    break
                merged = f"{merged.rstrip()} {lookahead.strip()}"
                consumed += 1
                if not has_stray_marker(merged):
                    stripped = merged
                    i += consumed
                    break

        blocks.append(("paragraph", stripped))
        i += 1

    return blocks


def renumber(blocks):
    """Resequence ordered-list numbers per nesting level.

    The markdown restarts numbering wherever the original conversion hit
    an intervening paragraph. Word documents usually want one continuous
    sequence, so this walks the blocks and rewrites the literal numbers,
    resetting a level whenever a shallower level advances.
    """
    counters = {}
    out = []
    for kind, payload in blocks:
        if kind != "ordered":
            # Headings and rules start a new outline; running prose does not.
            if kind in ("heading", "rule"):
                counters.clear()
            out.append((kind, payload))
            continue
        level, _, text = payload
        counters[level] = counters.get(level, 0) + 1
        for deeper in [k for k in counters if k > level]:
            del counters[deeper]
        out.append((kind, (level, str(counters[level]), text)))
    return out


# --------------------------------------------------------------------------
# Inline markup -> runs
# --------------------------------------------------------------------------

def parse_inline(text):
    """Split text into (content, bold, italic, mark) tuples.

    Once a span's delimiters are matched, any leftover asterisks inside it
    are markup noise, not content -- the original .docx conversion emitted
    a handful of these (`***Other **Signatory***`) -- so they are dropped.
    Asterisks outside any span stay literal: they mean the line is
    malformed in a way this function should not guess at. has_stray_marker
    reports those so they can be fixed at the source.
    """
    runs = []
    for piece in INLINE_RE.split(text):
        if not piece:
            continue
        if piece.startswith("<mark>") and piece.endswith("</mark>"):
            inner = piece[len("<mark>"):-len("</mark>")]
            # A placeholder may itself be emphasized: <mark>**[Company]**</mark>
            for content, bold, italic, _ in parse_inline(inner) or [(inner, False, False, False)]:
                runs.append((content, bold, italic, True))
        elif piece.startswith("***") and piece.endswith("***") and len(piece) > 6:
            runs.append((piece[3:-3].replace("*", ""), True, True, False))
        elif piece.startswith("**") and piece.endswith("**") and len(piece) > 4:
            runs.append((piece[2:-2].replace("*", ""), True, False, False))
        elif piece.startswith("*") and piece.endswith("*") and len(piece) > 2:
            runs.append((piece[1:-1].replace("*", ""), False, True, False))
        else:
            runs.append((piece, False, False, False))
    return runs


def has_stray_marker(text):
    """True if the line leaves an asterisk outside any matched span."""
    return any("*" in content and not (bold or italic)
               for content, bold, italic, _ in parse_inline(text))


def write_runs(paragraph, text, bold=False, italic=False):
    """Append text to a paragraph, honoring its inline markup."""
    for content, run_bold, run_italic, mark in parse_inline(text):
        if not content:
            continue
        run = paragraph.add_run(content)
        run.bold = bold or run_bold
        run.italic = italic or run_italic
        if mark:
            run.font.highlight_color = HIGHLIGHT
    return paragraph


# --------------------------------------------------------------------------
# Document assembly
# --------------------------------------------------------------------------

def set_base_style(doc):
    """Point the Normal style at the body font and give it legal spacing."""
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    # python-docx only sets the latin font; east-asian needs the raw property
    # or Word substitutes its own default for any non-ASCII character.
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    fmt = normal.paragraph_format
    fmt.space_after = Pt(8)
    fmt.line_spacing = 1.15

    for level, size in ((1, 14), (2, 12)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True


def set_page(doc):
    """US Letter, one-inch margins."""
    for section in doc.sections:
        section.start_type = WD_SECTION_START.NEW_PAGE
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        for side in ("top", "bottom", "left", "right"):
            setattr(section, f"{side}_margin", Inches(1))


def add_page_number_footer(doc):
    """Centered 'Page N of M' footer, built from Word field codes."""
    paragraph = doc.sections[0].footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def field(instruction):
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = instruction
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run = paragraph.add_run()._r
        for element in (begin, instr, end):
            run.append(element)

    paragraph.add_run("Page ")
    field(" PAGE ")
    paragraph.add_run(" of ")
    field(" NUMPAGES ")
    for run in paragraph.runs:
        run.font.name = BODY_FONT
        run.font.size = Pt(9)


def add_rule(doc):
    """A horizontal rule, drawn as a bottom border on an empty paragraph."""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    borders.append(bottom)
    paragraph._p.get_or_add_pPr().append(borders)


def add_table(doc, rows):
    """Render a markdown table as a bordered Word table with a bold header."""
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r, row in enumerate(rows):
        for c in range(width):
            cell = table.cell(r, c)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(2)
            write_runs(paragraph, row[c] if c < len(row) else "", bold=(r == 0))
    doc.add_paragraph()
    return table


def add_list_item(doc, text, level, marker):
    """A list item rendered as a paragraph with a hanging indent.

    The marker is literal text, not a Word list. See the module docstring.
    """
    paragraph = doc.add_paragraph()
    fmt = paragraph.paragraph_format
    fmt.left_indent = Inches(0.5 + 0.5 * level)
    fmt.first_line_indent = Inches(-0.5)
    fmt.space_after = Pt(6)
    paragraph.add_run(f"{marker}\t")
    write_runs(paragraph, text)
    return paragraph


def render(blocks, title=None):
    """Assemble parsed blocks into a Document."""
    doc = Document()
    set_base_style(doc)
    set_page(doc)
    add_page_number_footer(doc)

    doc.core_properties.title = title or ""
    doc.core_properties.author = "General Legal"
    doc.core_properties.comments = (
        "Generated from markdown at https://github.com/General-Legal/legal-templates. "
        "Released under CC0 1.0. Highlighted text marks fields to customize. "
        "Not legal advice."
    )

    for kind, payload in blocks:
        if kind == "heading":
            level, text = payload
            heading = doc.add_heading(level=min(level, 4))
            heading.text = ""
            write_runs(heading, text)
        elif kind == "rule":
            add_rule(doc)
        elif kind == "table":
            add_table(doc, payload)
        elif kind == "ordered":
            level, number, text = payload
            add_list_item(doc, text, level, f"{number}.")
        elif kind == "bullet":
            level, text = payload
            add_list_item(doc, text, level, "\u2022")
        else:
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            write_runs(paragraph, payload)

    return doc


def build(md, title=None, resequence=False):
    """Compile a markdown string into a Document."""
    blocks = parse_blocks(md)
    if resequence:
        blocks = renumber(blocks)
    return render(blocks, title=title)


# --------------------------------------------------------------------------
# Verification
# --------------------------------------------------------------------------

def highlighted_text(path):
    """Concatenate every highlighted run in a .docx, body and tables."""
    doc = Document(path)
    out = []

    def scan(paragraphs):
        for paragraph in paragraphs:
            for run in paragraph.runs:
                color = run.font.highlight_color
                if color is not None and color != WD_COLOR_INDEX.AUTO:
                    out.append(run.text)

    scan(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                scan(cell.paragraphs)
    return "".join(out)


def marked_text(md):
    """Concatenate the contents of every <mark> span, markup stripped."""
    spans = re.findall(r"<mark>(.*?)</mark>", md, re.DOTALL)
    return "".join(re.sub(r"\*+", "", s) for s in spans)


def stray_markers(blocks):
    """Snippets whose asterisks could not be resolved into emphasis.

    These are malformed in template.md itself -- usually a span opened in
    one table cell and closed in another. The build still succeeds and
    renders them literally; fixing them belongs in the markdown.
    """
    found = []

    def scan(text):
        if text and has_stray_marker(text):
            snippet = " ".join(text.split())
            found.append(snippet[:90] + ("..." if len(snippet) > 90 else ""))

    for kind, payload in blocks:
        if kind == "heading":
            scan(payload[1])
        elif kind == "table":
            for row in payload:
                for cell in row:
                    scan(cell)
        elif kind in ("ordered", "bullet"):
            scan(payload[-1])
        elif kind == "paragraph":
            scan(payload)
    return found


def check(md, path):
    """Confirm the placeholders survived the round trip.

    Compares the text inside <mark> spans with the highlighted text in
    the .docx, ignoring whitespace. Returns an error string or None.
    """
    expected = re.sub(r"\s+", "", marked_text(md))
    actual = re.sub(r"\s+", "", highlighted_text(path))
    if expected == actual:
        return None
    return (
        f"placeholder mismatch: {len(expected)} highlighted chars in markdown, "
        f"{len(actual)} in .docx"
    )


# --------------------------------------------------------------------------

def title_for(slug, template_dir):
    """Use the template README's H1 as the document title, if there is one."""
    readme = template_dir / "README.md"
    if readme.exists():
        for line in readme.read_text(encoding="utf-8").splitlines():
            if line.startswith("# "):
                return line[2:].strip()
    return slug.replace("-", " ").title()


def main():
    root = Path(__file__).resolve().parent.parent
    parser = argparse.ArgumentParser(description=__doc__,
                                     formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("slugs", nargs="*",
                        help="template directory names; omit to build all")
    parser.add_argument("--out", default="build/docx",
                        help="output directory (default: build/docx)")
    parser.add_argument("--renumber", action="store_true",
                        help="resequence ordered-list numbers instead of "
                             "copying them from the markdown")
    parser.add_argument("--zip", dest="make_zip", action="store_true",
                        help="also write <out>/../legal-templates-docx.zip")
    parser.add_argument("--no-check", dest="check", action="store_false",
                        help="skip the placeholder round-trip verification")
    args = parser.parse_args()

    templates = root / "templates"
    available = sorted(p.name for p in templates.iterdir()
                       if (p / "template.md").exists())

    selected = args.slugs or available
    unknown = [s for s in selected if s not in available]
    if unknown:
        parser.error(f"unknown template(s): {', '.join(unknown)}\n"
                     f"available: {', '.join(available)}")

    out_dir = (root / args.out) if not Path(args.out).is_absolute() else Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)

    failures = []
    warnings = []
    for slug in selected:
        template_dir = templates / slug
        md = (template_dir / "template.md").read_text(encoding="utf-8")

        doc = build(md, title=title_for(slug, template_dir),
                    resequence=args.renumber)

        dest = out_dir / f"{slug}.docx"
        doc.save(dest)

        note = ""
        if args.check:
            problem = check(md, dest)
            if problem:
                failures.append(f"{slug}: {problem}")
                note = f"  !! {problem}"
        strays = stray_markers(parse_blocks(md))
        if strays:
            warnings.append((slug, strays))
            note += f"  ({len(strays)} unresolved emphasis marker(s))"
        size = dest.stat().st_size
        print(f"{slug:<32} -> {dest.relative_to(root) if dest.is_relative_to(root) else dest} "
              f"({size:,} bytes){note}")

    if args.make_zip:
        archive = shutil.make_archive(str(out_dir.parent / "legal-templates-docx"),
                                      "zip", root_dir=out_dir)
        print(f"\nzip: {Path(archive).name} "
              f"({Path(archive).stat().st_size:,} bytes)")

    print(f"\nBuilt {len(selected)} template(s) into {out_dir}")

    if warnings:
        # Not a build failure: the text is rendered literally and nothing
        # is lost. It means template.md has an unbalanced * that a human
        # should close.
        print("\nUnresolved emphasis markers (fix these in template.md):")
        for slug, strays in warnings:
            for snippet in strays:
                print(f"  {slug}: {snippet}")

    if failures:
        print("\nPlaceholder check failed:", file=sys.stderr)
        for failure in failures:
            print(f"  {failure}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
